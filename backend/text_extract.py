"""
Text extraction from PDF, DOCX, and plain text.
All parsing is in-memory — no temp files written to disk.
"""
from __future__ import annotations
import io
import zipfile

_MAX_PAGES = 20
_MAX_TEXT_CHARS = 2_000_000  # 2 MB of text is already enormous for a CV

# Magic bytes for content-type verification (client Content-Type header is attacker-controlled)
_MAGIC = {
    ".pdf": b"%PDF",
    ".docx": b"PK\x03\x04",  # DOCX is a ZIP archive
}


def extract_text(raw: bytes, filename: str) -> str:
    """
    Extract visible + hidden text from raw file bytes.
    Raises ValueError with a user-facing message on parse failure.
    """
    lower = filename.lower()

    # Verify magic bytes regardless of the declared Content-Type
    for ext, magic in _MAGIC.items():
        if lower.endswith(ext):
            if not raw.startswith(magic):
                raise ValueError(
                    f"File does not appear to be a valid {ext.lstrip('.')} "
                    "(magic bytes mismatch). Upload a real PDF, DOCX, or TXT."
                )

    if lower.endswith(".pdf"):
        return _extract_pdf(raw)
    if lower.endswith(".docx"):
        return _extract_docx(raw)
    return raw.decode("utf-8", errors="replace")


# ── PDF ───────────────────────────────────────────────────────────────────────

def _extract_pdf(raw: bytes) -> str:
    visible = _extract_pdf_pdfminer(raw)
    hidden = _extract_pdf_hidden_text(raw)
    meta = _extract_pdf_metadata(raw)
    parts = [visible]
    if hidden.strip():
        parts.append(hidden)
    if meta.strip():
        parts.append(meta)
    return "\n".join(parts)


def _extract_pdf_metadata(raw: bytes) -> str:
    """User-controllable document info (Title/Author/Subject/Keywords). An injection
    can hide here just as in DOCX core properties. Tool-set fields (Producer, dates)
    are skipped to avoid adding noise to a clean CV. Best-effort."""
    try:
        import pdfplumber
    except ImportError:
        return ""
    try:
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            meta = pdf.metadata or {}
    except Exception:
        return ""
    return "\n".join(
        str(meta[k]) for k in ("Title", "Author", "Subject", "Keywords") if meta.get(k)
    )


def _extract_pdf_pdfminer(raw: bytes) -> str:
    try:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LAParams, LTTextContainer
    except ImportError as e:
        raise ValueError("PDF extraction library unavailable.") from e

    try:
        buf = io.BytesIO(raw)
        parts: list[str] = []
        for i, page_layout in enumerate(extract_pages(buf, laparams=LAParams())):
            if i >= _MAX_PAGES:
                break
            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    parts.append(element.get_text())
        text = "".join(parts)
        return text[:_MAX_TEXT_CHARS]
    except Exception as exc:
        raise ValueError(f"Could not parse PDF: {exc}") from exc


def _extract_pdf_hidden_text(raw: bytes) -> str:
    """
    pdfplumber layer: extract characters with font size < 3pt or white/near-white
    color that would be invisible in a rendered view.
    """
    try:
        import pdfplumber
    except ImportError:
        return ""

    chars: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for i, page in enumerate(pdf.pages):
                if i >= _MAX_PAGES:
                    break
                for ch in page.chars:
                    size = float(ch.get("size") or 12)
                    color = ch.get("non_stroking_color")
                    if size < 3.0 or _is_near_white(color):
                        chars.append(ch.get("text", ""))
    except Exception:
        return ""

    return "".join(chars)


def _is_near_white(color: object) -> bool:
    if color is None:
        return False
    if isinstance(color, (int, float)):
        return float(color) >= 0.9          # grayscale 0=black 1=white
    if isinstance(color, (tuple, list)):
        if len(color) == 3:                 # RGB
            return all(float(c) >= 0.9 for c in color)
        if len(color) == 4:                 # CMYK: (0,0,0,0) = white
            c, m, y, k = (float(x) for x in color)
            return c <= 0.1 and m <= 0.1 and y <= 0.1 and k <= 0.1
    return False


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _extract_docx(raw: bytes) -> str:
    try:
        import docx
        from docx.oxml.ns import qn
    except ImportError as e:
        raise ValueError("DOCX extraction library unavailable.") from e

    # Zip-bomb / structure preflight: inspect the central directory (no decompression)
    # before python-docx loads the archive, so a small file that inflates to gigabytes
    # cannot exhaust memory.
    try:
        zf = zipfile.ZipFile(io.BytesIO(raw))
    except zipfile.BadZipFile as exc:
        raise ValueError(f"Could not parse DOCX: {exc}") from exc
    with zf:
        infos = zf.infolist()
        names = {info.filename for info in infos}
        total_uncompressed = sum(info.file_size for info in infos)
    if len(infos) > 2000:
        raise ValueError("DOCX has too many internal entries.")
    if not {"[Content_Types].xml", "word/document.xml"}.issubset(names):
        raise ValueError("File is not a valid DOCX (missing OOXML parts).")
    if total_uncompressed > 50 * 1024 * 1024:
        raise ValueError("DOCX uncompressed size exceeds 50 MB limit.")

    try:
        doc = docx.Document(io.BytesIO(raw))
    except Exception as exc:
        raise ValueError(f"Could not parse DOCX: {exc}") from exc

    parts: list[str] = []

    # Body paragraphs (visible text)
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)

    # Headers and footers — white text or tiny text in headers is invisible to readers
    try:
        for section in doc.sections:
            for para in section.header.paragraphs:
                if para.text:
                    parts.append(para.text)
            for para in section.footer.paragraphs:
                if para.text:
                    parts.append(para.text)
    except Exception:  # nosec B110 -- optional DOCX section; missing is expected
        pass

    # TextBoxes (w:txbxContent — floating boxes, not in doc.paragraphs)
    try:
        for txbx in doc.element.body.iter(qn("w:txbxContent")):
            text = "".join(
                t.text for t in txbx.iter(qn("w:t")) if t.text
            )
            if text:
                parts.append(text)
    except Exception:  # nosec B110 -- optional DOCX section; missing is expected
        pass

    # Comments (/word/comments.xml relationship)
    try:
        for rel in doc.part.rels.values():
            if "comments" in rel.reltype.lower():
                for comment_el in rel.target_part._element.iter(qn("w:comment")):
                    text = "".join(
                        t.text for t in comment_el.iter(qn("w:t")) if t.text
                    )
                    if text:
                        parts.append(text)
    except Exception:  # nosec B110 -- optional DOCX section; missing is expected
        pass

    # Core document properties (title, subject, description, keywords)
    try:
        cp = doc.core_properties
        for attr in ("title", "subject", "description", "keywords",
                     "author", "last_modified_by", "comments"):
            val = getattr(cp, attr, None)
            if val:
                parts.append(str(val))
    except Exception:  # nosec B110 -- optional DOCX section; missing is expected
        pass

    text = "\n".join(parts)
    return text[:_MAX_TEXT_CHARS]
